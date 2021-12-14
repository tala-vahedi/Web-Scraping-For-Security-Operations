# Script Purpose: Process a target website and extract key information and suspicious content
# Script Version: 1.0 
# Script Author:  Tala Vahedi, University of Arizona

# Script Revision History:
# Version 1.0 Sept 15, 2021, Python 3.x

# 3rd Party Modules
from bs4 import BeautifulSoup 
from docx import Document  

# Python Standard Library
import requests
import os, glob
import re

# Psuedo Constants
SCRIPT_NAME    = "Script: Process a target website and extract key information and suspicious content"
SCRIPT_VERSION = "Version 1.0"
SCRIPT_AUTHOR  = "Author: Tala Vahedi"

if __name__ == '__main__':
    # Print Basic Script Information
    print()
    print(SCRIPT_NAME)
    print(SCRIPT_VERSION)
    print(SCRIPT_AUTHOR)
    print()  

    # creating a word document to capture all website contents
    document = Document()
    # adding word doc title
    document.add_heading('CYBV 473 Assignment 9 Output', 0)

    url = 'https://casl.website/login'
    base = 'https://casl.website'


    print("Sending request to ", str(url), " please wait...")
    # sending get request to the url
    page = requests.get(url)   
    # parsing through the text with bs4    
    soup = BeautifulSoup(page.text, 'html.parser')
    print("Parsing ", str(url), " with bs4, please wait...")
    # extracting the title of the website and converting to string format
    title = soup.title.string
    # adding the title of the website to the word doc
    document.add_paragraph('The title of the website is: ' + title + "\n")

    # adding the links found on the website to the word document
    document.add_paragraph('Links found on website: ') 
    # iterating through the website contents for href/links 
    for link in soup.findAll('a', attrs={'href': re.compile("^http://")}):
        # adding each link found in website to the word doc
        document.add_paragraph(link.get('href'))
    # entering a new line in word doc for aesthetic purposes
    document.add_paragraph("\n")

    print("Extracting Images from: ", str(url), " please wait...")

    # finding all image tags in website contents
    images = soup.findAll('img')  
    # adding the images on the website to the word doc
    document.add_paragraph("Images found on Website: ")
    # Process and display each image
    for eachImage in images:      
        try:
            # get source of each image
            imgURL = eachImage['src']
            # if URL path is relative
            if imgURL[0:4] != 'http':       
                # try prepending the base url
                imgURL = base+imgURL         
            # get the image from the URL
            response = requests.get(imgURL) 
            # get image name          
            imageName = os.path.basename(imgURL)
            # downloading and saving each image to current directory
            with open(imageName, 'wb') as outFile:
                outFile.write(response.content)
        # throwing an exception and continueing if error occurs
        except Exception as err:
            print(imgURL, err)
            continue   
    # for each png in current directory add img to word doc
    for img in glob.glob("*.png"):
        document.add_picture(img)
    # save the word doc
    document.save('vahediT_WK6_assignment9.docx') 
    print('\nScript Complete')